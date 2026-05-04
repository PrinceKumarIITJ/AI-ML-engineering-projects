import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def markdown_to_docx(md_path, docx_path):
    doc = Document()
    
    # Title styling
    style = doc.styles['Title']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(24)
    font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
        
    lines = content.split('\n')
    
    in_code_block = False
    
    for line in lines:
        if line.startswith('```'):
            in_code_block = not in_code_block
            if in_code_block:
                doc.add_paragraph("[Diagram/Code section starts]")
            else:
                doc.add_paragraph("[Diagram/Code section ends]")
            continue
            
        if in_code_block:
            p = doc.add_paragraph(line)
            p.style = 'No Spacing'
            if p.runs:
                p.runs[0].font.name = 'Courier New'
            continue
            
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=3)
        elif line.startswith('---'):
            pass # page break or horizontal rule, skip for simplicity
        elif line.startswith('* ') or line.startswith('- '):
            # Clean up bold markers
            clean_line = line[2:].replace('**', '')
            p = doc.add_paragraph(clean_line, style='List Bullet')
        elif line.startswith('1. ') or line.startswith('2. ') or line.startswith('3. ') or line.startswith('4. ') or line.startswith('5. '):
            clean_line = line[3:].replace('**', '')
            p = doc.add_paragraph(clean_line, style='List Number')
        elif line.strip() == '':
            pass
        else:
            # Clean up bold markers for standard paragraphs
            clean_line = line.replace('**', '')
            doc.add_paragraph(clean_line)

    doc.save(docx_path)
    print(f"Successfully saved {docx_path}")

if __name__ == "__main__":
    markdown_to_docx("MANAGER_PROJECT_REPORT.md", "IntelliLead_Project_Report.docx")
